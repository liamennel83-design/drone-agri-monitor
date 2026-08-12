#!/usr/bin/env python3
"""Génère le rapport Word du projet à partir des éléments du dépôt.

Les tableaux, le texte et les schémas construits avec des tableaux Word restent
éditables dans Word. Les figures PNG sont intégrées directement et leurs sources
sont listées en annexe et dans reports/editable_sources/.
"""

from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "reports" / "Rapport_Projet_Drone_Stress_Hydrique.docx"
FIGURES = ROOT / "reports" / "figures"

PRIMARY = "1F4E79"
ACCENT = "2E75B6"
LIGHT = "D9EAF7"
GRAY = "F2F2F2"
YELLOW = "FFF2CC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Page ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    paragraph.add_run(" sur ")
    fld2 = OxmlElement("w:fldSimple")
    fld2.set(qn("w:instr"), "NUMPAGES")
    paragraph._p.append(fld2)


def add_toc(document: Document) -> None:
    paragraph = document.add_paragraph()
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    paragraph._p.append(field)
    note = document.add_paragraph()
    note.add_run("Dans Word, cliquez avec le bouton droit sur cette table puis choisissez « Mettre à jour les champs ». ").italic = True


def add_title(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_text(document: Document, text: str, bold_start: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    if bold_start and text.startswith(bold_start):
        paragraph.add_run(bold_start).bold = True
        paragraph.add_run(text[len(bold_start):])
    else:
        paragraph.add_run(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_shading(header_cells[index], PRIMARY)
        set_cell_text(header_cells[index], header, bold=True, color="FFFFFF")
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], str(value))
            if len(table.rows) % 2 == 1:
                set_cell_shading(cells[index], GRAY)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    document.add_paragraph()
    return table


def add_figure(document: Document, filename: str, caption: str, width_cm: float = 15.0) -> None:
    path = FIGURES / filename
    if not path.exists():
        document.add_paragraph(f"Figure indisponible : {filename}")
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Cm(width_cm))
    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.style = "Caption"
    caption_p.add_run(caption)


def add_editable_flow(document: Document, labels: list[str]) -> None:
    table = document.add_table(rows=1, cols=len(labels) * 2 - 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for index, label in enumerate(labels):
        box = table.cell(0, index * 2)
        set_cell_shading(box, LIGHT)
        set_cell_text(box, label, bold=True, color=PRIMARY)
        if index < len(labels) - 1:
            arrow = table.cell(0, index * 2 + 1)
            set_cell_text(arrow, ">", bold=True, color=PRIMARY)
    document.add_paragraph("Schéma modifiable : cliquez dans chaque bloc pour modifier le texte.").italic = True


def add_placeholder(document: Document, label: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(label)
    run.bold = True
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    paragraph.paragraph_format.space_after = Pt(6)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in [("Title", 28, PRIMARY), ("Heading 1", 17, PRIMARY), ("Heading 2", 13, ACCENT), ("Heading 3", 11, PRIMARY)]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("Projet de drone autonome | Binôme A")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("666666")
    add_page_number(section.footer.paragraphs[0])
    props = document.core_properties
    props.title = "Rapport de projet - Drone autonome et stress hydrique"
    props.subject = "Rapport technique et expérimental"
    props.author = "Binôme A"
    props.comments = "Généré depuis le dépôt du projet."


def build_report() -> Document:
    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Cm(4.0)
    title.add_run("RAPPORT DE PROJET").bold = True
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.color.rgb = RGBColor.from_string(ACCENT)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Drone autonome pour le suivi du stress hydrique des plants de tomate")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor.from_string(PRIMARY)
    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Binôme A\nProjet de robot aérien autonome\nAntananarivo, août 2026").font.size = Pt(14)
    document.add_paragraph()
    add_placeholder(document, "À compléter avant envoi : noms des étudiants, établissement, encadrants, date officielle de remise.")
    document.add_page_break()

    add_title(document, "Résumé", 1)
    add_text(document, "Ce projet étudie l'utilisation d'un micro-drone équipé d'une caméra OV2640 pour acquérir des images de plants de tomate et suivre des indicateurs visuels associés au stress hydrique. La démarche associe une simulation de trajectoire lawnmower, un prototype de vol basique, une chaîne de traitement d'images reposant sur ExG et GRVI, ainsi qu'un protocole de constitution d'une base annotée.")
    add_text(document, "Les résultats de simulation sont séparés des résultats expérimentaux. Le dataset actuellement stocké dans le dépôt est synthétique et sert à vérifier le fonctionnement logiciel. La validation scientifique du classificateur nécessite des images réelles, plusieurs prises de vue par plant et une annotation associée à l'humidité du sol.")
    add_title(document, "Mots clés", 2)
    add_text(document, "Drone, agriculture de précision, stress hydrique, OV2640, ExG, GRVI, Random Forest, ArUco, IMU, trajectoire lawnmower.")
    document.add_page_break()

    add_title(document, "Table des matières", 1)
    add_toc(document)
    document.add_page_break()

    add_title(document, "1. Contexte et objectifs", 1)
    add_text(document, "Le stress hydrique affecte la croissance des plantes et peut entraîner une baisse de rendement. Une observation visuelle répétée permet de détecter certains signes, mais elle est longue à effectuer manuellement. Le projet propose un dispositif expérimental compact qui combine acquisition aérienne, traitement d'image et suivi du substrat.")
    add_title(document, "1.1 Résultats attendus", 2)
    add_table(document, ["Résultat attendu", "Élément produit", "Niveau de validation"], [
        ["Modèle de trajectoire", "Simulation lawnmower et optimisation", "Simulation"],
        ["Prototype de vol", "Décollage, stationnaire et atterrissage sécurisés", "Essais à documenter"],
        ["Base d'images annotées", "Arborescence, métadonnées et fichier par plant", "Acquisition réelle en cours"],
        ["Rapport et recommandations", "Analyse, limites et suite Master", "Document présent"],
    ], [5.2, 7.0, 4.0])

    add_title(document, "2. Architecture générale du système", 1)
    add_editable_flow(document, ["Drone et caméra", "Acquisition", "Indices ExG et GRVI", "Classification", "Rapport et suivi"])
    add_text(document, "Le traitement est conçu comme une chaîne modulaire. Les images sont d'abord acquises et annotées. Quatre variables sont extraites : moyenne et écart type de ExG, moyenne et écart type de GRVI. Elles sont ensuite utilisées dans une analyse exploratoire avec Random Forest.")
    add_table(document, ["Composant", "Rôle", "État à la date du rapport"], [
        ["Pydrone ESP32-S3", "Plateforme de vol", "Prototype matériel"],
        ["Caméra OV2640", "Acquisition d'images", "Paramètres à confirmer en essai"],
        ["IMU", "Mesure d'attitude", "Utilisée pour la stabilisation"],
        ["Marqueurs ArUco", "Correction visuelle envisagée", "Proposition soumise aux encadrants"],
        ["Capteur d'humidité du sol", "Référence expérimentale du stress", "À étalonner et intégrer au suivi"],
    ], [4.2, 6.5, 5.5])

    add_title(document, "3. Dispositif expérimental", 1)
    add_text(document, "La parcelle retenue mesure 1 m par 2 m et accueille 12 pots selon une grille de trois colonnes et quatre rangées. Pour comparer un groupe témoin et un groupe test, deux parcelles ou deux ensembles comparables de 12 pots peuvent être utilisés. Les plants doivent être homogènes autant que possible : même espèce, variété, stade et substrat.")
    add_table(document, ["Paramètre", "Valeur", "Statut"], [
        ["Parcelle", "1 m x 2 m", "Retenu"],
        ["Disposition", "3 colonnes x 4 rangées", "Retenu"],
        ["Nombre de pots par parcelle", "12", "Retenu"],
        ["Altitude nominale de prise de vue", "0,70 m", "Calculée, à mesurer en essai"],
        ["Vitesse nominale", "0,20 m/s", "Simulée, à confirmer"],
        ["GSD estimé", "0,44 mm/pixel", "Calculé"],
    ], [7.2, 4.4, 4.6])
    add_figure(document, "emprise_camera.png", "Figure 1. Emprise théorique de la caméra à l'altitude étudiée.")

    add_title(document, "3.1 Protocole d'humidité du sol", 2)
    add_text(document, "Le stress doit être documenté par une mesure de l'humidité du substrat. Après étalonnage relatif du capteur entre un état sec et un état humidifié, le groupe témoin maintient une humidité de référence. Le groupe test est maintenu progressivement autour de 45 à 55 pour cent de cette référence, sans dessèchement irréversible.")
    add_table(document, ["Moment", "Mesure", "Donnée à enregistrer"], [
        ["Avant le protocole", "Étalonnage sec et humide", "Valeur brute et référence relative"],
        ["Chaque jour", "Mesure avant arrosage", "Humidité, volume d'eau, observation"],
        ["Chaque prise de vue", "Image d'un pot identifié", "Nom de fichier, plant_id, jour, humidité"],
    ], [4.2, 6.1, 5.9])

    add_title(document, "4. Simulation de trajectoire", 1)
    add_text(document, "La trajectoire de couverture suit un motif lawnmower, aussi appelé boustrophédon. Les lignes parallèles permettent de couvrir la zone avec un recouvrement frontal et latéral défini. Les paramètres sont produits par les scripts du dossier `simulation/`.")
    add_figure(document, "trajectoire_lawnmower.png", "Figure 2. Exemple de trajectoire lawnmower générée par le module de simulation.")
    add_figure(document, "profil_vitesse.png", "Figure 3. Profil de vitesse utilisé dans la simulation cinématique.")
    add_figure(document, "comparaison_modes.png", "Figure 4. Comparaison de modes de prise de vue issue de la simulation.")
    add_text(document, "Les valeurs de simulation doivent être présentées comme des estimations. Elles seront consolidées par des essais de vol court avec la configuration matérielle réellement embarquée.")

    add_title(document, "5. Traitement des images", 1)
    add_title(document, "5.1 Indices retenus", 2)
    add_table(document, ["Indice", "Formule", "Utilisation"], [
        ["ExG", "2G - R - B", "Mettre en évidence la dominance du vert"],
        ["GRVI", "(R - G) / (R + G)", "Comparer les canaux rouge et vert"],
    ], [3.0, 5.5, 7.7])
    add_text(document, "La formule GRVI ci-dessus correspond à l'implémentation de `imagerie/stress_detector.py`. Cette convention doit être utilisée de manière identique dans les scripts et dans le rapport. Changer le signe de la formule impose de recalibrer les seuils et de réentraîner tout modèle associé.")
    add_editable_flow(document, ["Image RGB", "Masque végétation", "4 variables", "Validation par plant"])
    add_figure(document, "carte_ndvi.png", "Figure 5. Illustration d'une carte d'indice de végétation présente dans le dépôt. Cette figure doit être décrite comme une visualisation et non comme une mesure NDVI calibrée.")

    add_title(document, "5.2 Classification et limites méthodologiques", 2)
    add_text(document, "Le classificateur Random Forest utilise quatre variables : ExG moyen, ExG écart type, GRVI moyen et GRVI écart type. Le fichier historique `train_classifier.py` génère des données synthétiques. Son score ne doit pas être interprété comme une performance sur des plants réels.")
    add_table(document, ["Élément", "Pratique à retenir", "Risque évité"], [
        ["Séparation des données", "Séparer les plants avec plant_id", "Même plant présent en apprentissage et test"],
        ["Métriques", "F1, accuracy, kappa et MCC", "Lecture unique d'un score isolé"],
        ["Données actuelles", "Démo synthétique seulement", "Conclusion terrain non justifiée"],
        ["Données futures", "Plusieurs vues et mesures par pot", "Absence de traçabilité expérimentale"],
    ], [4.0, 6.7, 5.5])
    add_figure(document, "importance_features.png", "Figure 6. Importance des variables pour le modèle exploratoire fourni dans le dépôt.")
    add_figure(document, "matrice_confusion.png", "Figure 7. Matrice de confusion issue du jeu de démonstration. Elle ne remplace pas une évaluation sur images réelles.")
    add_figure(document, "comparaison_algorithmes.png", "Figure 8. Comparaison d'algorithmes présente dans le dépôt. Les conditions exactes d'évaluation doivent être indiquées avant toute conclusion comparative.")

    add_title(document, "6. Prototype de vol et navigation", 1)
    add_text(document, "Le résultat minimal attendu est un prototype de vol basique, documenté par un décollage, un stationnaire et un atterrissage. Les informations de sécurité et les mesures réellement observées doivent être consignées dans une fiche d'essai et accompagnées d'une vidéo courte.")
    add_placeholder(document, "À compléter : date d'essai, masse du drone, batterie utilisée, durée de stationnaire, dérive observée et lien vers la vidéo.")
    add_title(document, "6.1 Proposition IMU + ArUco", 2)
    add_text(document, "La correction mécanique du centre de gravité a été testée selon le binôme. Une navigation hybride IMU + ArUco est proposée pour compenser les limites du GPS sur une petite parcelle. Cette proposition doit être validée par les encadrants avant d'être annoncée comme une solution mise en oeuvre.")
    add_editable_flow(document, ["IMU", "Filtre de fusion", "Correction ArUco", "Stationnaire ou trajectoire"])
    add_table(document, ["Essai", "Mesure à relever", "Critère indicatif"], [
        ["Détection marqueur", "Taux de détection selon hauteur et lumière", "Valeur mesurée"],
        ["Position statique", "Écart avec ruban", "Objectif inférieur à 5 cm"],
        ["Stationnaire", "Dérive pendant 30 secondes", "Valeur mesurée"],
        ["Perte de marqueur", "Réaction de sécurité", "Stationnaire ou atterrissage"],
    ], [4.2, 7.1, 4.9])

    add_title(document, "7. Gestion de l'énergie et sécurité", 1)
    add_text(document, "La courbe de décharge LiPo est utilisée comme référence indicative. Les seuils doivent être validés sous charge avec la batterie réelle, car tension, température et vieillissement influencent la mesure.")
    add_table(document, ["Niveau", "Seuil de référence", "Action"], [
        ["Normal", "Supérieur à 3,80 V", "Mission possible sous surveillance"],
        ["Attention", "Environ 3,70 V", "Préparer le retour"],
        ["Retour recommandé", "3,65 V", "Interrompre la mission"],
        ["Urgence", "3,50 V", "Atterrir dès que possible"],
    ], [4.5, 5.1, 6.6])
    add_figure(document, "courbe_decharge_lipo.png", "Figure 9. Courbe de décharge LiPo utilisée comme référence de travail. Une mesure expérimentale reste nécessaire.")

    add_title(document, "8. Discussion et recommandations pour le Master", 1)
    add_text(document, "Le projet fournit une base technique utile : simulation de couverture, chaîne d'extraction d'indices, structure de dataset et protocole de suivi. Les éléments qui demandent encore une validation expérimentale sont la répétabilité du vol, la qualité des acquisitions réelles et la généralisation du classificateur à des plants distincts.")
    add_bullets(document, [
        "Constituer une base réelle avec plusieurs jours d'acquisition, plusieurs vues par plant et une mesure d'humidité associée.",
        "Conserver une séparation stricte par plant pendant la validation du classificateur.",
        "Caractériser la précision de position et le débit réel de la chaîne IMU, caméra et ArUco.",
        "Évaluer une carte de stress lorsque les images sont suffisamment nombreuses et géoréférencées ou repérées par la parcelle.",
        "Étudier l'intégration d'autres mesures : humidité du sol, température, éclairage et données météo.",
    ])

    add_title(document, "9. Conclusion", 1)
    add_text(document, "Le projet a établi une chaîne cohérente entre trajectoire simulée, acquisition embarquée, traitement d'images et protocole de données. La priorité immédiate est de documenter un vol basique sûr et de commencer une acquisition réelle annotée. Les résultats issus du dataset synthétique doivent rester identifiés comme des résultats de démonstration. La navigation IMU + ArUco et la cartographie automatisée constituent des pistes structurantes pour une poursuite en Master.")

    document.add_page_break()
    add_title(document, "Annexe A. Fiche d'essai de vol", 1)
    add_table(document, ["Champ", "Valeur à compléter"], [
        ["Date et lieu", ""],
        ["Configuration matérielle", ""],
        ["Masse du drone", ""],
        ["Batterie et tension avant essai", ""],
        ["Altitude visée", ""],
        ["Durée de stationnaire", ""],
        ["Dérive observée", ""],
        ["Résultat", ""],
        ["Lien vidéo ou fichier", ""],
    ], [6.2, 10.0])

    add_title(document, "Annexe B. Éléments modifiables et sources", 1)
    add_text(document, "Le texte, les tableaux, les listes et les schémas réalisés sous forme de tableaux Word sont modifiables directement dans ce document. Les figures PNG sont intégrées dans Word et peuvent être sélectionnées, déplacées, redimensionnées ou remplacées.")
    add_text(document, "Pour modifier les courbes à la source, utiliser les scripts de simulation du dossier `simulation/`, les scripts d'imagerie du dossier `imagerie/` et les fichiers CSV du dossier `reports/editable_sources/`. Après modification, régénérer les figures puis exécuter `scripts/generate_report_docx.py`.")
    add_table(document, ["Figure", "Fichier source inséré", "Source modifiable"], [
        ["1", "reports/figures/emprise_camera.png", "Calculs et paramètres dans docs et simulation"],
        ["2", "reports/figures/trajectoire_lawnmower.png", "simulation/lawnmower_planner.py"],
        ["3", "reports/figures/profil_vitesse.png", "simulation/flight_simulator.py"],
        ["4", "reports/figures/comparaison_modes.png", "simulation/flight_simulator.py"],
        ["5", "reports/figures/carte_ndvi.png", "À régénérer après acquisition réelle"],
        ["6 à 8", "Figures de classification", "imagerie et données annotées"],
        ["9", "reports/figures/courbe_decharge_lipo.png", "reports/editable_sources/parametres_mission.csv"],
    ], [2.0, 7.0, 7.2])
    return document


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_report()
    document.save(OUTPUT)
    print(f"Rapport généré : {OUTPUT}")


if __name__ == "__main__":
    main()
